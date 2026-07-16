"""
生成真实文档解析测试文件（PDF/Word/Excel/图片）
用于评估 OCR、表格提取、文本提取的正确率

用法:
    python scripts/generate_benchmark_files.py

生成的文件在 data/benchmark/ 下，每个文件都有已知内容。
解析后 vs 已知内容的对比 = 解析正确率。
"""
import sys, os, json
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

BENCHMARK_DIR = "data/benchmark"
os.makedirs(BENCHMARK_DIR, exist_ok=True)

# ========================
# 1. 生成 PDF（文字 + 表格）
# ========================
def generate_pdf():
    print("[1/4] Generating PDF...")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        print("  Skipped (need reportlab: pip install reportlab)")
        return

    # 注册中文字体（宋体）
    font_path = "C:/Windows/Fonts/simsun.ttc"
    FONT_NAME = "SimSun"
    try:
        pdfmetrics.registerFont(TTFont(FONT_NAME, font_path))
        FONT_BOLD = FONT_NAME
        print(f"  Using Chinese font: {font_path}")
    except Exception:
        FONT_NAME = "Helvetica"
        FONT_BOLD = "Helvetica-Bold"
        print("  Chinese font not available, using Helvetica")

    path = os.path.join(BENCHMARK_DIR, "benchmark.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()

    # 覆盖所有样式使用中文字体
    for s in styles.byName.values():
        if hasattr(s, 'fontName') and 'Title' in s.name:
            s.fontName = FONT_BOLD
        elif hasattr(s, 'fontName'):
            s.fontName = FONT_NAME

    elements = []

    elements.append(Paragraph("文档解析测试报告", styles['Title']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("一、项目概况", styles['Heading2']))
    elements.append(Paragraph(
        "本项目旨在验证Agentic RAG系统的多格式文档解析能力。"
        "测试涵盖PDF文字提取、表格提取、OCR识别等核心功能。"
        "本文件包含已知的文字内容和表格数据，解析后与实际内容对比可计算解析准确率。",
        styles['Normal']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("二、关键参数配置", styles['Heading2']))
    table_data = [
        ['参数名称', '参数值', '说明'],
        ['learning_rate', '2e-5', 'SFT微调学习率'],
        ['batch_size', '8', '训练批次大小'],
        ['num_epochs', '3', '训练轮次'],
        ['warmup_ratio', '0.03', '学习率预热比例'],
        ['lora_rank', '8', 'LoRA秩'],
        ['lora_alpha', '16', 'LoRA缩放因子'],
    ]
    t = Table(table_data, colWidths=[120, 80, 200])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), FONT_BOLD),
        ('FONTNAME', (0, 1), (-1, -1), FONT_NAME),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("三、性能指标", styles['Heading2']))
    elements.append(Paragraph(
        "系统在标准测试集上的性能表现如下："
        "文档解析成功率为95%，平均解析耗时2.3秒/文件。"
        "OCR文字识别准确率达92%，表格提取成功率88%。"
        "检索召回率（Recall@5）为0.87，回答忠实度（Faithfulness）为0.91。",
        styles['Normal']))

    doc.build(elements)
    print(f"  -> {path} ({os.path.getsize(path)} bytes)")


# ========================
# 2. 生成 Word
# ========================
def generate_word():
    print("[2/4] Generating Word...")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.table import _Cell

    path = os.path.join(BENCHMARK_DIR, "benchmark.docx")
    doc = Document()
    doc.add_heading('Word文档解析测试', level=0)

    doc.add_heading('一、引言', level=1)
    doc.add_paragraph(
        '本文档用于测试Agentic RAG系统的Word文档解析能力。'
        '包含标题、段落、表格、列表等多种格式。'
        '解析后的内容应与原始内容完全一致。')

    doc.add_heading('二、模型配置对比表', level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['模型', '参数量', '显存占用', '推理速度']
    data = [
        ['Qwen-7B', '7B', '16GB', '30 tokens/s'],
        ['Qwen-14B', '14B', '32GB', '18 tokens/s'],
        ['Qwen-72B', '72B', '160GB', '5 tokens/s'],
        ['LLaMA-3-8B', '8B', '18GB', '28 tokens/s'],
        ['DeepSeek-67B', '67B', '140GB', '6 tokens/s'],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_text

    doc.add_heading('三、关键结论', level=1)
    doc.add_paragraph('经过对比测试，得出以下结论：')
    doc.add_paragraph('1. Qwen-7B在消费级显卡上即可运行，性价比最高。', style='List Bullet')
    doc.add_paragraph('2. Qwen-72B性能最强但需要多卡部署，适合离线批处理场景。', style='List Bullet')
    doc.add_paragraph('3. DeepSeek-67B在数学推理任务上表现突出。', style='List Bullet')
    doc.add_paragraph('4. 模型量化（INT8/INT4）可降低显存占用50-75%。', style='List Bullet')

    doc.add_paragraph(
        '最后，建议根据实际业务场景选择合适的模型规模和部署方案。'
        '对于实时对话场景推荐7B-14B量化模型，'
        '对于离线分析场景可使用72B或更大模型。')

    doc.save(path)
    print(f"  -> {path} ({os.path.getsize(path)} bytes)")


# ========================
# 3. 生成 Excel
# ========================
def generate_excel():
    print("[3/4] Generating Excel...")
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    path = os.path.join(BENCHMARK_DIR, "benchmark.xlsx")
    wb = openpyxl.Workbook()

    # Sheet 1: 季度销售数据
    ws1 = wb.active
    ws1.title = "季度销售"
    ws1.append(["季度", "产品", "销售额(万)", "同比增长", "环比增长"])
    sales_data = [
        ["Q1 2024", "A产品", 1280, "12.5%", "3.2%"],
        ["Q1 2024", "B产品", 856, "8.3%", "-1.5%"],
        ["Q1 2024", "C产品", 2340, "22.1%", "7.8%"],
        ["Q2 2024", "A产品", 1450, "15.2%", "13.3%"],
        ["Q2 2024", "B产品", 920, "9.8%", "7.5%"],
        ["Q2 2024", "C产品", 2680, "25.4%", "14.5%"],
    ]
    for row in sales_data:
        ws1.append(row)
    # 加粗表头
    for cell in ws1[1]:
        cell.font = Font(bold=True)

    # Sheet 2: 模型评测
    ws2 = wb.create_sheet("模型评测")
    ws2.append(["模型名称", "MMLU", "C-Eval", "HumanEval", "GSM8K"])
    eval_data = [
        ["Qwen-7B", 62.5, 72.3, 35.1, 61.2],
        ["Qwen-14B", 69.8, 78.2, 42.6, 68.5],
        ["Qwen-72B", 78.3, 84.5, 51.8, 76.4],
        ["LLaMA-3-8B", 65.3, 50.2, 38.2, 58.7],
        ["DeepSeek-67B", 72.1, 80.6, 47.3, 72.9],
    ]
    for row in eval_data:
        ws2.append(row)
    for cell in ws2[1]:
        cell.font = Font(bold=True)

    wb.save(path)
    print(f"  -> {path} ({os.path.getsize(path)} bytes)")


# ========================
# 4. 生成图片（含文字，用于OCR测试）
# ========================
def generate_image():
    print("[4/4] Generating image for OCR test...")
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("  Skipped (need Pillow)")
        return

    path = os.path.join(BENCHMARK_DIR, "benchmark_ocr.png")

    # 创建一个带文字的图片
    img = Image.new('RGB', (800, 500), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # 尝试加载中文字体
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",       # 微软雅黑
        "C:/Windows/Fonts/simhei.ttf",      # 黑体
        "C:/Windows/Fonts/simsun.ttc",      # 宋体
    ]
    font = None
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 24)
                break
            except Exception:
                continue

    if font is None:
        font = ImageFont.load_default()

    # 绘制文字内容（已知内容，用于验证OCR准确率）
    texts = [
        ("Agentic RAG OCR 测试", (150, 30), (0, 0, 0)),
        ("系统名称：多格式智能检索系统", (100, 80), (50, 50, 50)),
        ("版本号：v1.1.0", (100, 120), (50, 50, 50)),
        ("测试日期：2024年7月", (100, 160), (50, 50, 50)),
        ("关键指标：", (100, 220), (0, 0, 0)),
        ("  解析成功率：95%", (120, 260), (0, 80, 0)),
        ("  OCR识别率：92%", (120, 300), (0, 80, 0)),
        ("  表格提取率：88%", (120, 340), (0, 80, 0)),
        ("  向量检索维度：1536", (120, 380), (0, 0, 80)),
        ("  支持格式：PDF/Word/Excel/图片", (120, 420), (80, 0, 0)),
    ]

    for text, pos, color in texts:
        draw.text(pos, text, fill=color, font=font)

    # 画一个简单的表格
    draw.rectangle([100, 460, 700, 464], fill=(0, 0, 0))

    img.save(path)
    print(f"  -> {path} ({os.path.getsize(path)} bytes)")


# ========================
# 主入口
# ========================
if __name__ == "__main__":
    print("Generating benchmark files...\n")

    # 生成 manifest（记录每个文件的已知内容）
    manifest = {
        "generated_at": __import__('time').strftime('%Y-%m-%d %H:%M:%S'),
        "files": []
    }

    if len(sys.argv) > 1:
        # 只生成指定文件
        for arg in sys.argv[1:]:
            if arg == "pdf":
                generate_pdf()
                manifest["files"].append("benchmark.pdf")
            elif arg == "word":
                generate_word()
                manifest["files"].append("benchmark.docx")
            elif arg == "excel":
                generate_excel()
                manifest["files"].append("benchmark.xlsx")
            elif arg == "image":
                generate_image()
                manifest["files"].append("benchmark_ocr.png")
    else:
        generate_pdf()
        generate_word()
        generate_excel()
        generate_image()
        manifest["files"] = ["benchmark.pdf", "benchmark.docx", "benchmark.xlsx", "benchmark_ocr.png"]

    # 保存 manifest
    manifest_path = os.path.join(BENCHMARK_DIR, "manifest.json")
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    print(f"\nManifest saved: {manifest_path}")
    print("Done! Now run 'pytest tests/test_parser.py -v -k \"test_parser_stats\"'")

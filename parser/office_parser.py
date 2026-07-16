"""
Office 文档解析器 - Word/Excel
"""
import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger("parser.office")


class OfficeParser:
    def parse(self, file_path: str) -> List[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext == ".docx":
            return self._parse_docx(path)
        elif ext in (".xlsx", ".xls"):
            return self._parse_excel(path)
        else:
            raise ValueError(f"不支持的 Office 格式: {ext}")

    def _parse_docx(self, path: Path) -> List[Document]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            parts = []

            # 提取段落
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # 提取表格
            for ti, table in enumerate(doc.tables):
                parts.append(f"\n[表格 {ti+1}]")
                # 表头
                header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
                parts.append(f"| {header} |")
                parts.append("|" + " --- |" * len(table.rows[0].cells))
                # 数据行
                for row in table.rows[1:]:
                    row_data = " | ".join(cell.text.strip() for cell in row.cells)
                    parts.append(f"| {row_data} |")

            content = "\n".join(parts)
            return [Document(
                page_content=content,
                metadata={
                    "source": str(path),
                    "source_type": "docx",
                    "file_name": path.name,
                },
            )]
        except Exception as e:
            logger.error(f"Word 解析失败: {e}")
            return []

    def _parse_excel(self, path: Path) -> List[Document]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            docs = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join(
                        [str(c) if c is not None else "" for c in row]
                    )
                    if row_str.strip():
                        rows.append(row_str)
                content = f"=== 工作表: {sheet_name} ({len(rows)}行) ===\n" + "\n".join(rows)
                docs.append(Document(
                    page_content=content,
                    metadata={
                        "source": str(path),
                        "source_type": "xlsx" if path.suffix == ".xlsx" else "xls",
                        "file_name": path.name,
                        "sheet": sheet_name,
                    },
                ))
            wb.close()
            return docs
        except Exception as e:
            logger.error(f"Excel 解析失败: {e}")
            return []
